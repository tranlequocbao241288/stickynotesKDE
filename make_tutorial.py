import os
from docx import Document
from docx.shared import Pt, Inches

def create_tutorial_docx():
    doc = Document()
    
    # Tiêu đề chính
    title = doc.add_heading('GIÁO TRÌNH A-Z: LẬP TRÌNH WIDGET STICKY NOTES CHO KDE PLASMA', 0)
    title.alignment = 1 # Center

    doc.add_paragraph('Hướng dẫn chi tiết từ số 0 cho người mới bắt đầu lập trình, giải thích toàn bộ thuật toán, kiểu dữ liệu và cách xây dựng ứng dụng Sticky Notes hoàn chỉnh.')

    # ---------------------------------------------------------
    doc.add_heading('Phần 1: Giới thiệu chung & Kiến trúc dự án', level=1)
    doc.add_paragraph('Dự án chúng ta vừa xây dựng là một "Plasmoid" - một tiện ích nằm trên màn hình (Desktop Widget) của hệ điều hành Linux sử dụng môi trường KDE Plasma. Ứng dụng này cho phép người dùng tạo các tờ giấy nhớ (Sticky Notes) và danh sách công việc (Todo List) trực tiếp trên Desktop.')
    
    doc.add_heading('1.1. Công nghệ sử dụng', level=2)
    doc.add_paragraph('- QML (Qt Modeling Language): Một ngôn ngữ thiết kế giao diện theo dạng khai báo. QML được thiết kế tối ưu cho việc tạo ra giao diện người dùng trơn tru, dễ đọc, có cách viết khá giống JSON hoặc CSS.')
    doc.add_paragraph('- JavaScript (JS): Ngôn ngữ lập trình chịu trách nhiệm cho "bộ não" của ứng dụng. Mọi thuật toán lọc, tìm kiếm, lưu trữ dữ liệu đều được viết bằng JS.')
    
    doc.add_heading('1.2. Cấu trúc thư mục', level=2)
    code = """package/
 ├── metadata.json           (File cấu hình cơ bản của widget)
 └── contents/
      ├── code/
      │    └── logic.js      (Chứa mọi giải thuật, xử lý dữ liệu)
      └── ui/
           ├── main.qml      (Giao diện chính, màn hình bọc ngoài)
           ├── NoteCard.qml  (Giao diện của một tờ ghi chú)
           ├── TodoItem.qml  (Giao diện của một dòng công việc)
           └── TodoInput.qml (Khung nhập công việc mới)"""
    p = doc.add_paragraph(code)
    p.style = 'No Spacing'
    p.paragraph_format.left_indent = Inches(0.5)

    # ---------------------------------------------------------
    doc.add_heading('Phần 2: Nền tảng Lập trình (Dành cho người chưa biết gì)', level=1)
    
    doc.add_heading('2.1. Biến (Variables) & Kiểu dữ liệu (Data Types)', level=2)
    doc.add_paragraph('Trong lập trình, "Biến" giống như những chiếc hộp dùng để cất giữ thông tin. Mỗi hộp sẽ chứa một loại đồ vật nhất định, gọi là "Kiểu dữ liệu". Trong dự án này, ta sử dụng các kiểu dữ liệu sau:')
    doc.add_paragraph('1. String (Chuỗi chữ): Dùng để lưu văn bản. Ví dụ: Tiêu đề note "Công việc cần làm", màu sắc "#fff59d".')
    doc.add_paragraph('2. Number (Số thẳng): Lưu giá trị số, ví dụ vị trí X, Y (100, 150) hoặc chiều rộng/cao của Note.')
    doc.add_paragraph('3. Boolean (Đúng/Sai): Chỉ có 2 giá trị true (đúng) hoặc false (sai). Ví dụ trạng thái hoàn thành công việc: completed = true.')
    
    doc.add_heading('Khái niệm nâng cao: Array và Object', level=3)
    doc.add_paragraph('4. Array (Mảng): Một danh sách các giá trị nằm cạnh nhau. Trong ứng dụng, danh sách toàn bộ các Note hoặc danh sách Todo trong mỗi note là một mảng.')
    doc.add_paragraph('Ví dụ mảng: [ "Note 1", "Note 2", "Note 3" ]')
    doc.add_paragraph('5. Object (Đối tượng): Một hộp chứa nhiều ngăn, mỗi ngăn có một nhãn (key) và giá trị (value). Để biểu diễn 1 Note, ta không thể dùng một chuỗi đơn giản, mà dùng Object:')
    code = """var myNote = {
    id: "a1b2-c3d4",
    title: "Mua sắm",
    color: "#fff59d",
    items: [] // Đây là mảng chứa các Todo
};"""
    doc.add_paragraph(code).paragraph_format.left_indent = Inches(0.5)

    doc.add_heading('2.2. Hàm (Function) & Tham số (Parameters)', level=2)
    doc.add_paragraph('Hàm là một xưởng máy. Bạn đưa nguyên liệu (Tham số - Parameters) vào, xưởng máy xử lý và trả ra thành phẩm (Return).')
    doc.add_paragraph('Ví dụ: Hàm tạo công việc mới (createTodo). Đưa vào tham số là "noteId", kết quả trả ra là một Object Todo trống (chứa ID ngẫu nhiên, ngày tạo là hôm nay, nội dung rỗng).')

    # ---------------------------------------------------------
    doc.add_heading('Phần 3: Các Giải thuật và Logic Cốt lõi (Trong logic.js)', level=1)
    
    doc.add_heading('3.1. Thuật toán tạo ID ngẫu nhiên không trùng lặp (UUID)', level=2)
    doc.add_paragraph('Để quản lý Note và Todo, máy tính cần phân biệt chúng. Ta không dùng số thứ tự 1, 2, 3 vì nếu xóa số 2, thứ tự sẽ bị lệch. Giải pháp là thuật toán UUID v4 (chuỗi 36 ký tự ngẫu nhiên).')
    code = """function generateUUID() {
    return 'xxxxxxxx-xxxx-4xxx-yxxx-xxxxxxxxxxxx'.replace(/[xy]/g, function(c) {
        var r = Math.random() * 16 | 0;
        var v = c === 'x' ? r : (r & 0x3 | 0x8);
        return v.toString(16);
    });
}"""
    doc.add_paragraph(code).paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph('Cách hoạt động: Thuật toán tìm các ký tự x và y, rồi thay thế bằng một số ngẫu nhiên từ 0 đến 15 (hệ thập lục phân từ 0-9, a-f). Xác suất 2 ID trùng nhau khi sinh ngẫu nhiên là gần như bằng 0.')

    doc.add_heading('3.2. Thuật toán "Immutable Update" (Cập nhật bất biến)', level=2)
    doc.add_paragraph('Trong QML, khi danh sách thay đổi, giao diện sẽ chỉ cập nhật nếu nó thấy ĐỐI TƯỢNG MỚI được gán. Nếu ta sửa trực tiếp dữ liệu (Ví dụ: note.title = "A"), QML sẽ không nhận ra. Thuật toán ta sử dụng là .map() và Object.assign().')
    doc.add_paragraph('- Hàm map(): Đi qua từng phần tử trong mảng, tạo ra một mảng hoàn toàn mới.')
    doc.add_paragraph('- Hàm Object.assign({}, note, changes): Tạo một "clone" (Bản sao) của ghi chú, đồng thời ghi đè thông tin mới.')
    
    doc.add_heading('3.3. Thuật toán Lọc (Filter) và Tìm kiếm (Search)', level=2)
    doc.add_paragraph('Khi người dùng nhập từ khóa tìm kiếm (Ví dụ: "Họp"), thuật toán chạy như sau:')
    doc.add_paragraph('Bước 1: Chuyển toàn bộ dữ liệu và từ khóa về viết thường (toLowerCase) để so sánh không phân biệt Hoa/thường.')
    doc.add_paragraph('Bước 2: Tìm ở Title của Note: title.indexOf("họp"). Hàm indexOf tìm xem chữ "họp" có nằm trong Tiêu đề không.')
    doc.add_paragraph('Bước 3: Dùng vòng lặp quét tất cả các TodoItem bên trong Note đó. Nếu nội dung Todo có chứa "họp", lưu kết quả lại.')
    doc.add_paragraph('Bước 4: Dùng hàm filter() lọc ra danh sách các Note khớp dữ liệu, và kết hợp hàm map() để loại bỏ các todo không khớp nếu người dùng đang dùng nút Filter "Active" hoặc "Completed".')

    # ---------------------------------------------------------
    doc.add_heading('Phần 4: Thiết kế Giao diện với QML', level=1)
    
    doc.add_heading('4.1. Hệ tọa độ và Binding', level=2)
    doc.add_paragraph('QML vẽ giao diện theo tọa độ Oxy (x tính từ trái sang phải, y từ trên xuống dưới). Đặc điểm mạnh mẽ nhất của QML là "Binding" (Ràng buộc dữ liệu).')
    doc.add_paragraph('Ví dụ: Khi bạn kéo thả một ghi chú (NoteCard), tọa độ X và Y của Note liên tục thay đổi. Ta dùng DragHandler để bắt sự kiện chuột. Sự kiện onActiveChanged kiểm tra: Nếu chuột thả ra (active = false), gọi Javascript lưu tọa độ (x, y) này xuống Object Database. Nhờ Binding, khi lưu xong, Widget nhớ luôn vị trí trên màn hình.')

    doc.add_heading('4.2. Khắc phục lỗi Layout & Canvas (Kinh nghiệm thực tiễn)', level=2)
    doc.add_paragraph('Trong dự án, chúng ta dùng thẻ Flickable làm Canvas (Vùng không gian rộng để kéo Note). Tuy nhiên, có một lỗi nghiêm trọng: "Kéo Note bị di chuyển cả bảng".')
    doc.add_paragraph('Khắc phục: Ta tắt tính năng kéo-để-cuộn của Flickable (bằng lệnh "interactive: false"). Thay vào đó ta viết một chức năng WheelHandler để bắt con lăn chuột, cho phép người dùng cuộn canvas khi dùng thanh cuộn hoặc con lăn chuột, giúp việc tương tác với Note không bị kẹt.')
    
    doc.add_heading('4.3. Listview và Delegate', level=2)
    doc.add_paragraph('Ứng dụng hiện danh sách Todo bằng Repeater / ListView. Nguyên lý hoạt động là "Vòng lặp giao diện". Bạn đưa cho thẻ ListView một mảng 10 phần tử Todo, ListView tự động tạo ra 10 thẻ "TodoItem". Delegate chính là bản thiết kế của thẻ TodoItem đó.')
    doc.add_paragraph('Xử lý tràn chữ: Đối với văn bản nhập vào như Tiêu đề, phần mềm ban đầu xài TextField (chỉ hỗ trợ 1 dòng). Sau khi nâng cấp, ta dùng TextArea và bật lệnh "wrapMode: Text.Wrap". Lệnh này tự động tính toán chiều ngang, gặp hết dòng là gập text xuống mà không vượt qua giới hạn của Note.')

    # ---------------------------------------------------------
    doc.add_heading('Phần 5: Lưu Trữ Lâu Dài trên Hệ thống (Persistence)', level=1)
    doc.add_paragraph('Plasma KDE hỗ trợ lưu cấu hình Widget thông qua "plasmoid.configuration". Dữ liệu của chúng ta là Object/Array (Bộ nhớ RAM). Để lưu xuống Ổ cứng (Disk), ta phải ép chúng thành Chuỗi văn bản chung, gọi là chuỗi JSON.')
    doc.add_paragraph('Ta dùng hàm JSON.stringify() để nén toàn bộ biến Dữ liệu thành Text, lưu xuống hệ điều hành. Lần sau mở máy lên, Widget gọi JSON.parse() để biến Text trở về dạng Biến dữ liệu.')
    
    doc.add_heading('5.1 Debounce (Kỹ thuật tối ưu hiệu năng)', level=2)
    doc.add_paragraph('Nếu bạn gõ phím "T, e, s, t", ứng dụng sẽ phải Save vào Ổ cứng 4 lần. Rất nóng máy! Ta dùng kỹ thuật setTimeout (Hẹn giờ). Bạn gõ "T" -> Máy tính đếm 0.5 giây chờ đợi. Bạn gõ "e" lập tức -> Đồng hồ đếm 0.5s lại từ đầu. Cho đến khi bạn dừng gõ hẳn 0.5 giây, máy tính mới thực hiện lưu 1 LẦN. Đây gọi là kỹ thuật Debounce rất nổi tiếng trong lập trình.')

    # ---------------------------------------------------------
    doc.add_heading('Tổng Kết Lộ Trình Tự Thực Hành', level=1)
    doc.add_paragraph('Dành cho bạn muốn clone hoặc làm lại dự án này trên máy tính của mình:')
    doc.add_paragraph('Bước 1: Tạo file metadata.json để hệ điều hành nhận dạng widget của bạn.')
    doc.add_paragraph('Bước 2: Xây dựng File logic.js chứa tất cả hàm Khởi tạo Note, Sửa Note, Xóa Note, Format Tgian.')
    doc.add_paragraph('Bước 3: Viết main.qml làm nền, tạo thanh công cụ và tạo vùng kéo thả (Canvas). Load dữ liệu lên Repeater.')
    doc.add_paragraph('Bước 4: Thiết kế NoteCard.qml (Trang trí Màu sắc, Tiêu đề, gắn sự kiện kéo chuột). Tích hợp TodoItem.qml vào.')
    doc.add_paragraph('Bước 5: Thử nghiệm thực tế. Tích hợp lệnh Save và chạy kquitapp6 plasmashell && kstart plasmashell.')

    doc.add_paragraph('\n-- Hết Giáo Trình --')

    doc.save('Giao_trinh_StickyNotes_KDE.docx')
    print("Document successfully created: Giao_trinh_StickyNotes_KDE.docx")

if __name__ == "__main__":
    create_tutorial_docx()
